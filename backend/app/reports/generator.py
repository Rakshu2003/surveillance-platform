from docx import Document
from docx.shared import RGBColor
from datetime import datetime
import os

def generate_incident_report(events: list, video_id: str) -> str:
    doc = Document()

    title = doc.add_heading('SENTINELAI — INCIDENT REPORT', 0)
    title.runs[0].font.color.rgb = RGBColor(0, 100, 200)

    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Video ID: {video_id}')
    doc.add_paragraph(f'Total Events: {len(events)}')
    doc.add_paragraph(f'Critical (L3): {len([e for e in events if e.get("severity") == "L3"])}')
    doc.add_paragraph(f'Suspicious (L2): {len([e for e in events if e.get("severity","").startswith("L2")])}')
    doc.add_paragraph(f'Normal (L1): {len([e for e in events if e.get("severity") == "L1"])}')

    doc.add_heading('Event Details', level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Event ID'
    hdr[1].text = 'Type'
    hdr[2].text = 'Severity'
    hdr[3].text = 'Threat Score'
    hdr[4].text = 'Frame'

    for event in events[:50]:
        row = table.add_row().cells
        row[0].text = str(event.get('event_id', ''))
        row[1].text = str(event.get('event_type', ''))
        row[2].text = str(event.get('severity', ''))
        row[3].text = f"{event.get('threat_score', 0)*100:.1f}%"
        row[4].text = str(event.get('frame_index', ''))

    os.makedirs('reports', exist_ok=True)
    path = f'reports/incident_{video_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(path)
    return path
